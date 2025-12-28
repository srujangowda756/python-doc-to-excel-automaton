import pandas as pd
from docx import Document

doc = Document("input.docx")

def table_to_df(table):
    data = []
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])
    return data

first_table_data = table_to_df(doc.tables[0])
final_df = pd.DataFrame(first_table_data[1:], columns=first_table_data[0])

for i in range(1, len(doc.tables)):
    table_data = table_to_df(doc.tables[i])
    df = pd.DataFrame(table_data[1:], columns=final_df.columns)
    final_df = pd.concat([final_df, df], ignore_index=True)

final_df.to_excel("output.xlsx",index=False)